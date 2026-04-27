"""
Export report text to PDF and Word (.docx).
"""
import io
import logging
import re
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

DEFAULT_REPORT_TITLE = "IT License and Cost Optimization Report"
LEGACY_REPORT_TITLES = {
    "SQL Server License Optimization Report",
}
REPORT_LOGO_SVG_PATH = Path("static") / "assets" / "report_bayer_logo.svg"
REPORT_LOGO_PNG_PATH = Path("static") / "assets" / "report_bayer_logo.png"
TEXT_REPLACEMENTS = {
    "\r\n": "\n",
    "\r": "\n",
    "\u00a0": " ",
    "\u2192": "->",
    "\u2014": "-",
    "\u2013": "-",
}
EURO_SYMBOL = "\u20ac"
MONETARY_CONTEXT_LABELS = (
    "total estimated license cost",
    "total annual license cost",
    "total license cost",
    "total annual cost",
    "annual cost",
    "license cost",
    "license price",
    "average price",
    "avg price",
    "total cost",
    "cost savings",
    "expenses",
    "expense",
    "savings",
    "spend",
    "price",
    "cost",
    "money",
)
MONETARY_CONTEXT_PATTERN = re.compile(
    rf"(?i)(?P<label>\b(?:{'|'.join(re.escape(label) for label in MONETARY_CONTEXT_LABELS)})\b"
    r"(?:\s+(?:of|is|was|at))?"
    r"(?:\*{0,2})?\s*[:=-]?\s*(?:\*{0,2})?\s*)"
    rf"(?P<amount>(?<![{re.escape(EURO_SYMBOL)}$])\d[\d,]*(?:\.\d+)?)"
)


def _project_root() -> Path:
    try:
        from django.conf import settings

        return Path(settings.BASE_DIR)
    except Exception:
        return Path(__file__).resolve().parents[2]


def _get_asset_path(relative_path: Path) -> Optional[Path]:
    candidate = _project_root() / relative_path
    return candidate if candidate.exists() else None


def _normalize_report_text(text: str) -> str:
    normalized = text or ""
    for source, target in TEXT_REPLACEMENTS.items():
        normalized = normalized.replace(source, target)
    return normalized.strip()


def normalize_report_currency_text(text: str) -> str:
    normalized = (text or "").replace("$", EURO_SYMBOL)

    def _replace_amount(match: re.Match) -> str:
        return f"{match.group('label')}{EURO_SYMBOL}{match.group('amount')}"

    return MONETARY_CONTEXT_PATTERN.sub(_replace_amount, normalized)


def normalize_report_content_text(text: str) -> str:
    normalized = normalize_report_currency_text(_normalize_report_text(text))
    for legacy_title in LEGACY_REPORT_TITLES:
        normalized = re.sub(
            rf"(?mi)^#\s*{re.escape(legacy_title)}\s*$",
            f"# {DEFAULT_REPORT_TITLE}",
            normalized,
        )
        normalized = re.sub(
            rf"(?mi)^{re.escape(legacy_title)}\s*$",
            DEFAULT_REPORT_TITLE,
            normalized,
        )
    return normalized


def normalize_report_title_text(text: str) -> str:
    return normalize_report_content_text(text)


def _ordinal(value: int) -> str:
    if 10 <= value % 100 <= 20:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(value % 10, "th")
    return f"{value}{suffix}"


def _resolve_generated_at(generated_at=None):
    if generated_at is not None:
        return generated_at
    try:
        from django.utils import timezone

        return timezone.localtime()
    except Exception:
        return datetime.now()


def _format_generated_at(generated_at=None) -> str:
    stamp = _resolve_generated_at(generated_at)
    return f"{_ordinal(stamp.day)} {stamp.strftime('%B %Y | %I:%M %p')}"


def _find_font_file(directory: Path, patterns) -> Optional[Path]:
    if not directory.exists():
        return None
    for pattern in patterns:
        matches = sorted(directory.glob(pattern))
        if matches:
            return matches[0]
    return None


def _resolve_pdf_fonts():
    fonts = {"regular": "Helvetica", "bold": "Helvetica-Bold"}
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return fonts

    font_dirs = [
        _project_root() / "static" / "assets" / "fonts",
        Path(r"C:\Windows\Fonts"),
    ]
    regular_patterns = [
        "Acumin*Regular*.ttf",
        "Acumin*Regular*.otf",
        "Acumin*Book*.ttf",
        "Acumin*Book*.otf",
    ]
    bold_patterns = [
        "Acumin*Semibold*.ttf",
        "Acumin*Semibold*.otf",
        "Acumin*Bold*.ttf",
        "Acumin*Bold*.otf",
    ]

    regular_path = None
    bold_path = None
    for font_dir in font_dirs:
        regular_path = regular_path or _find_font_file(font_dir, regular_patterns)
        bold_path = bold_path or _find_font_file(font_dir, bold_patterns)

    if regular_path:
        try:
            pdfmetrics.registerFont(TTFont("AcuminProRegular", str(regular_path)))
            fonts["regular"] = "AcuminProRegular"
        except Exception:
            logger.warning("Unable to register Acumin Pro regular font; falling back to Helvetica")
    if bold_path:
        try:
            pdfmetrics.registerFont(TTFont("AcuminProBold", str(bold_path)))
            fonts["bold"] = "AcuminProBold"
        except Exception:
            logger.warning("Unable to register Acumin Pro bold font; falling back to Helvetica-Bold")

    return fonts


def _draw_bayer_logo(canvas, x: float, y: float, size: float, fonts):
    outer_pad = size * 0.08
    inner_pad = size * 0.13

    canvas.saveState()
    canvas.setLineCap(1)
    canvas.setLineWidth(size * 0.08)

    outer = canvas.beginPath()
    outer.arc(x + outer_pad, y + outer_pad, x + size - outer_pad, y + size - outer_pad, 40, 200)
    canvas.setStrokeColorRGB(137 / 255.0, 211 / 255.0, 41 / 255.0)
    canvas.drawPath(outer, stroke=1, fill=0)

    inner = canvas.beginPath()
    inner.arc(x + inner_pad, y + inner_pad, x + size - inner_pad, y + size - inner_pad, 220, 200)
    canvas.setStrokeColorRGB(0 / 255.0, 188 / 255.0, 255 / 255.0)
    canvas.drawPath(inner, stroke=1, fill=0)

    canvas.setFillColorRGB(16 / 255.0, 56 / 255.0, 79 / 255.0)
    canvas.setFont(fonts["bold"], size * 0.195)
    canvas.drawCentredString(x + size * 0.5, y + size * 0.475, "BAYER")

    small_size = size * 0.155
    canvas.setFont(fonts["bold"], small_size)
    canvas.drawCentredString(x + size * 0.5, y + size * 0.80, "B")
    canvas.drawCentredString(x + size * 0.5, y + size * 0.66, "A")
    canvas.drawCentredString(x + size * 0.5, y + size * 0.28, "E")
    canvas.drawCentredString(x + size * 0.5, y + size * 0.14, "R")
    canvas.restoreState()


def _markdown_to_plain(text: str) -> str:
    """Strip markdown markup for plain text outputs."""
    text = normalize_report_title_text(text)
    text = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def _markdown_to_reportlab(text: str) -> str:
    safe = escape(normalize_report_title_text(text))
    safe = re.sub(r"`([^`]+)`", r"<font name='Courier'>\1</font>", safe)
    safe = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", safe)
    safe = re.sub(r"\*([^*]+)\*", r"<i>\1</i>", safe)
    return safe


def _append_paragraph_block(blocks, lines):
    text = " ".join(line.strip() for line in lines if line.strip())
    if text:
        blocks.append({"kind": "paragraph", "text": text})


def _parse_table_row(line: str) -> list[str]:
    """Extract cell values from a markdown table row like | col1 | col2 |."""
    return [cell.strip() for cell in line.strip("|").split("|")]


def _is_table_separator(line: str) -> bool:
    """Return True for markdown table separator rows like |---|---:|."""
    return bool(re.fullmatch(r"\|[-: |]+\|", line))


def _parse_report_blocks(report_text: str):
    report_text = normalize_report_title_text(report_text)
    title = None
    blocks = []
    paragraph_lines = []

    def flush_paragraph():
        nonlocal paragraph_lines
        _append_paragraph_block(blocks, paragraph_lines)
        paragraph_lines = []

    for raw_line in report_text.splitlines():
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue
        if re.fullmatch(r"-{3,}", line):
            flush_paragraph()
            blocks.append({"kind": "rule"})
            continue
        if line.startswith("# "):
            flush_paragraph()
            title = line[2:].strip() or title
            continue
        if line.startswith("## "):
            flush_paragraph()
            blocks.append({"kind": "section", "text": line[3:].strip()})
            continue
        if line.startswith("### "):
            flush_paragraph()
            blocks.append({"kind": "subsection", "text": line[4:].strip()})
            continue
        if re.match(r"^[-*]\s+", line):
            flush_paragraph()
            blocks.append({"kind": "bullet", "text": re.sub(r"^[-*]\s+", "", line)})
            continue
        if re.match(r"^\d+\.\s+", line):
            flush_paragraph()
            match = re.match(r"^(\d+\.)\s+(.*)$", line)
            if match:
                blocks.append({"kind": "numbered", "label": match.group(1), "text": match.group(2)})
            continue
        # Markdown table rows
        if line.startswith("|") and line.endswith("|"):
            flush_paragraph()
            if _is_table_separator(line):
                continue  # skip separator lines
            cells = _parse_table_row(line)
            blocks.append({"kind": "table_row", "cells": cells})
            continue
        if re.match(r"^[A-Z][A-Za-z0-9/&() ,' -]{0,80}:$", line):
            flush_paragraph()
            blocks.append({"kind": "subsection", "text": line})
            continue
        paragraph_lines.append(line)

    flush_paragraph()
    return title or DEFAULT_REPORT_TITLE, blocks


def _to_int(value) -> int:
    try:
        return int(value or 0)
    except (TypeError, ValueError):
        return 0


def _to_float(value) -> float:
    try:
        return float(value or 0)
    except (TypeError, ValueError):
        return 0.0


def _format_count(value) -> str:
    return f"{_to_int(value):,}"


def format_currency(value) -> str:
    return f"{EURO_SYMBOL}{_to_float(value):,.2f}"


def _format_currency(value) -> str:
    return format_currency(value)


def _default_executive_summary(report_context) -> str:
    total_demand = _format_count(report_context.get("total_demand_quantity"))
    total_cost = _format_currency(report_context.get("total_license_cost"))
    azure_count = _format_count(report_context.get("azure_payg_count"))
    retired_count = _format_count(report_context.get("retired_count"))
    return (
        "This report evaluates the current state of IT licensing and identifies key opportunities for cost "
        f"optimization. The organization manages {total_demand} licenses at a total annual cost of {total_cost}. "
        f"Key opportunities include transitioning {azure_count} devices from Azure BYOL to PAYG and addressing "
        f"{retired_count} retired devices with active installations. These actions can reduce costs, improve "
        "compliance, and enhance data quality."
    )


def _extract_executive_summary(report_text: str, report_context) -> str:
    _, blocks = _parse_report_blocks(report_text)
    current_section = None
    summary_paragraphs = []
    for block in blocks:
        if block["kind"] == "section":
            current_section = block["text"].strip().lower()
            continue
        if current_section == "executive summary" and block["kind"] == "paragraph":
            summary_paragraphs.append(_markdown_to_plain(block["text"]))
    if summary_paragraphs:
        return " ".join(summary_paragraphs[:2])
    return _default_executive_summary(report_context)


def _resolve_savings_value(report_context, direct_key: str, rule_key: Optional[str] = None):
    value = report_context.get(direct_key)
    if value is None and rule_key:
        value = (report_context.get("rule_wise_savings") or {}).get(rule_key)
    return value


def _build_template_blocks(report_text: str, report_context):
    azure_count = _format_count(report_context.get("azure_payg_count"))
    retired_count = _format_count(report_context.get("retired_count"))
    total_demand = _format_count(report_context.get("total_demand_quantity"))
    total_cost = _format_currency(report_context.get("total_license_cost"))
    total_savings = _format_currency(_resolve_savings_value(report_context, "total_savings"))
    azure_payg_savings = _format_currency(_resolve_savings_value(report_context, "azure_payg_savings", "azure_payg"))
    retired_devices_savings = _format_currency(
        _resolve_savings_value(report_context, "retired_devices_savings", "retired_devices")
    )
    summary = _extract_executive_summary(report_text, report_context)

    blocks = [
        {"kind": "section", "text": "Executive Summary"},
        {"kind": "paragraph", "text": summary},
        {"kind": "section", "text": "Current State"},
        {"kind": "subsection", "text": "License Demand"},
        {"kind": "bullet", "text": f"Total licenses in use: {total_demand}"},
        {
            "kind": "bullet",
            "text": (
                f"Azure BYOL licenses: Significant subset, with {azure_count} devices identified as potential "
                "PAYG candidates."
            ),
        },
        {
            "kind": "bullet",
            "text": (
                f"Retired devices: {retired_count} devices with active installations, indicating potential "
                "data quality issues."
            ),
        },
        {"kind": "subsection", "text": "Cost"},
        {"kind": "bullet", "text": f"Total annual license cost: {total_cost}"},
        {
            "kind": "bullet",
            "text": "High costs associated with underutilized Azure BYOL licenses and licenses on retired devices.",
        },
        {"kind": "subsection", "text": "Savings"},
        {"kind": "bullet", "text": f"Total savings: {total_savings}"},
        {"kind": "bullet", "text": f"BYOL to PAYG Savings: {azure_payg_savings}"},
        {"kind": "bullet", "text": f"Retired but reporting Savings: {retired_devices_savings}"},
        {"kind": "subsection", "text": "Product Mix"},
        {
            "kind": "bullet",
            "text": "Majority of licenses are tied to cloud services, with Azure being a significant cost driver.",
        },
        {
            "kind": "bullet",
            "text": "Legacy licenses still active on retired devices, contributing to inefficiencies.",
        },
        {"kind": "section", "text": "Optimization Opportunities"},
        {"kind": "subsection", "text": f"Azure BYOL -> PAYG ({azure_count} Devices)"},
        {"kind": "bullet_group", "text": "Benefits:"},
        {
            "kind": "sub_bullet",
            "text": "PAYG (Pay-As-You-Go) eliminates upfront costs and aligns expenses with actual usage.",
        },
        {"kind": "sub_bullet", "text": "Reduces over-provisioning risks and improves cost predictability."},
        {
            "kind": "sub_bullet",
            "text": "Simplifies license management by removing the need for manual tracking of BYOL compliance.",
        },
        {"kind": "bullet_group", "text": "Risks:"},
        {"kind": "sub_bullet", "text": "Potential cost increases if usage patterns are not monitored and optimized."},
        {"kind": "sub_bullet", "text": "Requires accurate usage forecasting to avoid unexpected expenses."},
        {"kind": "subsection", "text": f"Retired Devices ({retired_count} Devices)"},
        {"kind": "bullet_group", "text": "Benefits:"},
        {
            "kind": "sub_bullet",
            "text": "Decommissioning licenses on retired devices reduces waste and ensures compliance.",
        },
        {"kind": "sub_bullet", "text": "Improves data quality by removing outdated records from asset inventories."},
        {"kind": "bullet_group", "text": "Risks:"},
        {
            "kind": "sub_bullet",
            "text": "Incomplete or inaccurate asset tracking may lead to missed decommissioning opportunities.",
        },
        {
            "kind": "sub_bullet",
            "text": "Potential compliance risks if retired devices are still linked to active contracts.",
        },
        {"kind": "section", "text": "Risks"},
        {"kind": "subsection", "text": "Data Quality:"},
        {"kind": "bullet", "text": "Inaccurate asset inventories may lead to over-licensing or under-licensing."},
        {
            "kind": "bullet",
            "text": "Retired devices with active installations create discrepancies in usage reporting.",
        },
        {"kind": "subsection", "text": "Compliance:"},
        {"kind": "bullet", "text": "Mismanagement of BYOL licenses can result in non-compliance with vendor agreements."},
        {"kind": "bullet", "text": "Retired devices pose risks of unauthorized usage or audit penalties."},
        {"kind": "subsection", "text": "Cost:"},
        {"kind": "bullet", "text": "Failure to optimize Azure licensing may lead to continued overspending."},
        {"kind": "bullet", "text": "PAYG transition requires careful monitoring to avoid unexpected cost spikes."},
        {"kind": "section", "text": "Recommendations"},
        {"kind": "subsection", "text": f"Transition Azure BYOL to PAYG for {azure_count} Devices:"},
        {"kind": "bullet", "text": "Conduct a detailed usage analysis to confirm PAYG suitability."},
        {"kind": "bullet", "text": "Implement monitoring tools to track usage and costs post-transition."},
        {"kind": "subsection", "text": "Decommission Licenses on Retired Devices:"},
        {"kind": "bullet", "text": f"Audit the {retired_count} retired devices to confirm decommissioning eligibility."},
        {"kind": "bullet", "text": "Update asset inventories to reflect accurate device and license statuses."},
        {"kind": "subsection", "text": "Improve Data Quality:"},
        {"kind": "bullet", "text": "Implement regular asset inventory reviews to ensure accuracy."},
        {"kind": "bullet", "text": "Leverage automated tools to track device lifecycle and license usage."},
        {"kind": "subsection", "text": "Monitor PAYG Usage Post-Transition:"},
        {"kind": "bullet", "text": "Set up alerts for unusual usage patterns to prevent cost overruns."},
        {
            "kind": "bullet",
            "text": "Regularly review PAYG expenses to identify further optimization opportunities.",
        },
        {"kind": "subsection", "text": "Enhance Compliance Processes:"},
        {"kind": "bullet", "text": "Establish clear policies for managing BYOL and PAYG licenses."},
        {"kind": "bullet", "text": "Train IT staff on compliance requirements and best practices."},
    ]
    return DEFAULT_REPORT_TITLE, blocks


def _get_render_blocks(report_text: str, report_context=None):
    if report_context:
        return _build_template_blocks(report_text, report_context)
    return _parse_report_blocks(report_text)


def build_report_markdown(report_text: str, report_context=None) -> str:
    report_title, blocks = _get_render_blocks(report_text, report_context=report_context)
    lines = [f"# {report_title}"]

    for block in blocks:
        kind = block["kind"]
        text = block.get("text", "")
        if kind == "section":
            lines.extend(["", f"## {text}"])
        elif kind == "subsection":
            lines.extend(["", f"### {text}"])
        elif kind == "bullet":
            lines.append(f"- {text}")
        elif kind == "bullet_group":
            lines.append(f"- {text}")
        elif kind == "sub_bullet":
            lines.append(f"  - {text}")
        elif kind == "numbered":
            lines.append(f"{block.get('label', '1.')} {text}")
        elif kind == "rule":
            lines.extend(["", "---"])
        elif kind == "table_row":
            cells = block.get("cells") or []
            lines.append("| " + " | ".join(cells) + " |")
        else:
            lines.extend(["", text])

    return normalize_report_content_text("\n".join(lines).strip())


def export_pdf(report_text: str, generated_at=None, report_context=None) -> Optional[bytes]:
    """Generate branded PDF bytes from report text. Returns None if reportlab not installed."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        logger.warning("reportlab not installed; PDF export disabled")
        return None

    fonts = _resolve_pdf_fonts()
    report_title, blocks = _get_render_blocks(report_text, report_context=report_context)
    generated_label = _format_generated_at(generated_at)
    svg_logo_path = _get_asset_path(REPORT_LOGO_SVG_PATH)
    png_logo_path = _get_asset_path(REPORT_LOGO_PNG_PATH)
    svg_logo = None
    render_pdf = None

    if svg_logo_path:
        try:
            from reportlab.graphics import renderPDF
            from svglib.svglib import svg2rlg

            svg_logo = svg2rlg(str(svg_logo_path))
            render_pdf = renderPDF
        except ImportError:
            logger.warning("svglib not installed; falling back to PNG logo for PDF export")
        except Exception:
            logger.exception("Unable to load SVG logo for report header; falling back to PNG")
            svg_logo = None
            render_pdf = None

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        topMargin=1.62 * inch,
        bottomMargin=0.8 * inch,
    )

    palette = {
        "title": colors.HexColor("#10384F"),
        "heading": colors.HexColor("#000000"),
        "text": colors.HexColor("#111111"),
        "muted": colors.HexColor("#7D8691"),
        "rule": colors.HexColor("#D3DDE5"),
    }

    styles = getSampleStyleSheet()
    section_style = ParagraphStyle(
        name="ReportSection",
        parent=styles["Heading2"],
        fontName=fonts["bold"],
        fontSize=15.5,
        leading=18,
        spaceBefore=10,
        spaceAfter=3,
        textColor=palette["heading"],
    )
    subsection_style = ParagraphStyle(
        name="ReportSubsection",
        parent=styles["Heading3"],
        fontName=fonts["bold"],
        fontSize=12.2,
        leading=14,
        spaceBefore=5,
        spaceAfter=2,
        textColor=palette["muted"],
    )
    body_style = ParagraphStyle(
        name="ReportBody",
        parent=styles["BodyText"],
        fontName=fonts["regular"],
        fontSize=10.95,
        leading=13.7,
        textColor=palette["text"],
        spaceAfter=3,
    )
    bullet_style = ParagraphStyle(
        name="ReportBullet",
        parent=body_style,
        leftIndent=18,
        firstLineIndent=0,
        bulletIndent=0,
        spaceAfter=2,
    )
    bullet_group_style = ParagraphStyle(
        name="ReportBulletGroup",
        parent=body_style,
        leftIndent=18,
        firstLineIndent=0,
        bulletIndent=0,
        spaceAfter=1,
    )
    sub_bullet_style = ParagraphStyle(
        name="ReportSubBullet",
        parent=body_style,
        leftIndent=38,
        firstLineIndent=0,
        bulletIndent=20,
        spaceAfter=1,
    )

    story = [Spacer(1, 0.02 * inch)]
    for block in blocks:
        kind = block["kind"]
        if kind == "section":
            story.append(Paragraph(_markdown_to_reportlab(block["text"]), section_style))
        elif kind == "subsection":
            story.append(Paragraph(_markdown_to_reportlab(block["text"]), subsection_style))
        elif kind == "bullet":
            story.append(Paragraph(_markdown_to_reportlab(block["text"]), bullet_style, bulletText="\u2022"))
        elif kind == "bullet_group":
            story.append(Paragraph(_markdown_to_reportlab(block["text"]), bullet_group_style, bulletText="\u2022"))
        elif kind == "sub_bullet":
            story.append(Paragraph(_markdown_to_reportlab(block["text"]), sub_bullet_style, bulletText="\u2022"))
        elif kind == "numbered":
            story.append(Paragraph(_markdown_to_reportlab(block["text"]), bullet_style, bulletText=block["label"]))
        elif kind == "rule":
            story.append(HRFlowable(width="100%", thickness=0.6, color=palette["rule"], spaceBefore=8, spaceAfter=8))
        elif kind == "table_row":
            cells = block.get("cells") or []
            row_text = "  |  ".join(_markdown_to_reportlab(c) for c in cells if c)
            story.append(Paragraph(row_text, bullet_style))
        else:
            story.append(Paragraph(_markdown_to_reportlab(block["text"]), body_style))

    def draw_page(canvas, pdf_doc):
        page_width, page_height = A4
        left = pdf_doc.leftMargin
        right = page_width - pdf_doc.rightMargin
        logo_size = 46
        logo_x = left + 2
        logo_y = page_height - 67

        canvas.saveState()
        canvas.setStrokeColor(palette["rule"])
        canvas.setLineWidth(0.7)
        canvas.line(0.45 * inch, page_height - 14, page_width - 0.45 * inch, page_height - 14)
        if svg_logo and render_pdf:
            scale = min(logo_size / max(svg_logo.width, 1), logo_size / max(svg_logo.height, 1))
            canvas.saveState()
            canvas.translate(logo_x, logo_y)
            canvas.scale(scale, scale)
            render_pdf.draw(svg_logo, canvas, 0, 0)
            canvas.restoreState()
        elif png_logo_path:
            canvas.drawImage(
                str(png_logo_path),
                logo_x,
                logo_y,
                width=logo_size,
                height=logo_size,
                preserveAspectRatio=True,
                mask="auto",
            )
        else:
            _draw_bayer_logo(canvas, logo_x, logo_y, logo_size, fonts)

        title_x = left + 58
        canvas.setFillColor(palette["title"])
        canvas.setFont(fonts["bold"], 18.2)
        canvas.drawString(title_x, page_height - 43, report_title)

        canvas.setFillColor(palette["muted"])
        canvas.setFont(fonts["regular"], 11)
        canvas.drawString(left + 1, page_height - 91, generated_label)

        canvas.setFont(fonts["regular"], 10)
        canvas.drawRightString(right, 38, f"Page | {canvas.getPageNumber()}")
        canvas.restoreState()

    doc.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    return buf.getvalue()


def export_docx(report_text: str, generated_at=None, report_context=None) -> Optional[bytes]:
    """Generate Word document bytes. Returns None if python-docx not installed."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
    except ImportError:
        logger.warning("python-docx not installed; Word export disabled")
        return None

    report_title, blocks = _get_render_blocks(report_text, report_context=report_context)
    generated_label = _format_generated_at(generated_at)
    logo_path = _get_asset_path(REPORT_LOGO_PNG_PATH)

    doc = Document()
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Title"):
        try:
            doc.styles[style_name].font.name = "Acumin Pro"
        except Exception:
            pass

    if logo_path:
        doc.add_picture(str(logo_path), width=Inches(0.62))

    title = doc.add_heading(report_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(generated_label)

    for block in blocks:
        kind = block["kind"]
        text = _markdown_to_plain(block.get("text", ""))
        if kind == "section":
            doc.add_heading(text, level=1)
        elif kind == "subsection":
            doc.add_heading(text, level=2)
        elif kind in {"bullet", "bullet_group", "sub_bullet"}:
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "numbered":
            doc.add_paragraph(text, style="List Number")
        elif kind == "rule":
            doc.add_paragraph("")
        elif kind == "table_row":
            cells = block.get("cells") or []
            row_text = "  |  ".join(_markdown_to_plain(c) for c in cells if c)
            if row_text:
                doc.add_paragraph(row_text)
        elif text:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def export_xlsx(report_text: str, generated_at=None, report_context=None) -> Optional[bytes]:
    """Generate Excel workbook bytes. Returns None if openpyxl not installed."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    except ImportError:
        logger.warning("openpyxl not installed; Excel export disabled")
        return None

    report_title, blocks = _get_render_blocks(report_text, report_context=report_context)
    generated_label = _format_generated_at(generated_at)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Report"
    sheet.sheet_view.showGridLines = False
    sheet.freeze_panes = "A4"

    title_fill = PatternFill(fill_type="solid", fgColor="10384F")
    section_fill = PatternFill(fill_type="solid", fgColor="EAF4FB")
    subsection_fill = PatternFill(fill_type="solid", fgColor="F8FAFC")
    thin_rule = Side(style="thin", color="D3DDE5")
    wrap_alignment = Alignment(vertical="top", wrap_text=True)

    sheet.merge_cells("A1:D1")
    title_cell = sheet["A1"]
    title_cell.value = report_title
    title_cell.font = Font(name="Calibri", size=16, bold=True, color="FFFFFF")
    title_cell.fill = title_fill
    title_cell.alignment = Alignment(horizontal="left", vertical="center")

    sheet.merge_cells("A2:D2")
    meta_cell = sheet["A2"]
    meta_cell.value = generated_label
    meta_cell.font = Font(name="Calibri", size=10, italic=True, color="5B6470")
    meta_cell.alignment = Alignment(horizontal="left", vertical="center")

    current_row = 4
    section_row = None

    for block in blocks:
        kind = block["kind"]
        raw_text = _markdown_to_plain(block.get("text", ""))

        if kind == "section":
            current_row += 1
            sheet.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=4)
            cell = sheet.cell(row=current_row, column=1, value=raw_text)
            cell.font = Font(name="Calibri", size=13, bold=True, color="10384F")
            cell.fill = section_fill
            cell.alignment = wrap_alignment
            cell.border = Border(top=thin_rule, bottom=thin_rule)
            section_row = current_row
        elif kind == "subsection":
            current_row += 1
            sheet.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=4)
            cell = sheet.cell(row=current_row, column=1, value=raw_text)
            cell.font = Font(name="Calibri", size=11, bold=True, color="475569")
            cell.fill = subsection_fill
            cell.alignment = wrap_alignment
        elif kind in {"bullet", "bullet_group", "sub_bullet", "numbered", "paragraph"}:
            current_row += 1
            label = ""
            if kind == "sub_bullet":
                label = "  - "
            elif kind == "numbered":
                label = f"{block.get('label', '1.')} "
            elif kind in {"bullet", "bullet_group"}:
                label = "- "
            cell = sheet.cell(row=current_row, column=1, value=f"{label}{raw_text}")
            cell.font = Font(name="Calibri", size=10)
            cell.alignment = wrap_alignment
        elif kind == "table_row":
            current_row += 1
            cells = block.get("cells") or []
            for col_idx, cell_value in enumerate(cells[:4], start=1):
                cell = sheet.cell(row=current_row, column=col_idx, value=_markdown_to_plain(cell_value))
                cell.font = Font(name="Calibri", size=10)
                cell.alignment = wrap_alignment
        elif kind == "rule":
            current_row += 1
            for col in range(1, 5):
                cell = sheet.cell(row=current_row, column=col)
                cell.border = Border(bottom=thin_rule)
        else:
            current_row += 1
            cell = sheet.cell(row=current_row, column=1, value=raw_text)
            cell.font = Font(name="Calibri", size=10)
            cell.alignment = wrap_alignment

        if section_row and current_row > section_row:
            for col in range(1, 5):
                sheet.cell(row=current_row, column=col).alignment = wrap_alignment

    sheet.column_dimensions["A"].width = 110
    sheet.column_dimensions["B"].width = 18
    sheet.column_dimensions["C"].width = 18
    sheet.column_dimensions["D"].width = 18

    for row_idx in range(1, current_row + 1):
        sheet.row_dimensions[row_idx].height = 22

    buf = io.BytesIO()
    workbook.save(buf)
    buf.seek(0)
    return buf.getvalue()
